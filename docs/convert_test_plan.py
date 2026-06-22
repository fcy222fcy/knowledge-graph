# -*- coding: utf-8 -*-
"""
将软件测试计划文档markdown转换为docx格式
格式要求：
- 所有英文：Times New Roman
- 一级标题：三号宋体(16pt)、加粗
- 二级标题：四号宋体(14pt)、加粗
- 三级标题：小四宋体(12pt)、加粗
- 正文：小四宋体(12pt)
- 行间距：固定值20磅
- 段落首行缩进2个字符
- 表格内文字：五号宋体(10.5pt)，表头加粗
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', font_size=None, bold=False):
    """设置字体"""
    run.font.name = font_name_en
    run.font.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)


def set_paragraph_format(paragraph, line_spacing=20, first_line_indent=True):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Cm(0.85)


def parse_markdown_to_docx(md_file, docx_file):
    """解析markdown并生成docx"""
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表格缓冲
    in_table = False
    table_data = []

    def flush_table():
        nonlocal table_data, in_table
        if table_data:
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            if rows > 0 and cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row_data):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text.strip())
                        set_font(run, font_size=10.5, bold=(r_idx == 0))
                        set_paragraph_format(p, first_line_indent=False)
            table_data = []
            in_table = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            if in_table:
                flush_table()
            i += 1
            continue

        # 跳过分隔线
        if stripped in ('---', '***'):
            if in_table:
                flush_table()
            i += 1
            continue

        # 表格行
        if stripped.startswith('|') and stripped.endswith('|'):
            # 检查是否为分隔行
            inner = stripped[1:-1]
            if all(c in '-| ' for c in inner):
                i += 1
                continue
            cells = [c.strip() for c in inner.split('|')]
            if not in_table:
                in_table = True
            table_data.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # 一级标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, font_size=16, bold=True)
            set_paragraph_format(p, first_line_indent=False)
            i += 1
            continue

        # 二级标题
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_font(run, font_size=14, bold=True)
            set_paragraph_format(p, first_line_indent=False)
            i += 1
            continue

        # 三级标题
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_font(run, font_size=12, bold=True)
            set_paragraph_format(p, first_line_indent=False)
            i += 1
            continue

        # 四级标题
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_font(run, font_size=12, bold=True)
            set_paragraph_format(p, first_line_indent=False)
            i += 1
            continue

        # 列表项（- 开头）
        if stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # 添加缩进
            pf = p.paragraph_format
            pf.left_indent = Cm(0.85)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(20)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            # 处理列表项中的加粗文本
            run = p.add_run('• ')
            set_font(run, font_size=12, bold=False)
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part[2:-2]
                    run = p.add_run(bold_text)
                    set_font(run, font_size=12, bold=True)
                else:
                    if part:
                        run = p.add_run(part)
                        set_font(run, font_size=12, bold=False)
            i += 1
            continue

        # 正文（处理加粗文本）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_format(p, first_line_indent=True)

        # 分割并处理加粗文本
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 加粗文本
                text = part[2:-2]
                run = p.add_run(text)
                set_font(run, font_size=12, bold=True)
            else:
                # 普通文本
                if part:
                    run = p.add_run(part)
                    set_font(run, font_size=12, bold=False)
        i += 1

    if in_table:
        flush_table()

    doc.save(docx_file)
    print(f'文档已保存到: {docx_file}')


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(base_dir, '软件测试计划文档.md')
    docx_file = os.path.join(base_dir, '软件测试计划文档.docx')
    parse_markdown_to_docx(md_file, docx_file)
