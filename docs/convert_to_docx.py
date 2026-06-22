# -*- coding: utf-8 -*-
"""
将第7章软件测试markdown文档转换为符合格式要求的docx文档
格式要求：
- 所有英文：Times New Roman
- 一级标题：三号宋体(16pt)、加粗
- 二级标题：四号宋体(14pt)、加粗
- 三级标题：小四宋体(12pt)、加粗
- 正文：小四宋体(12pt)
- 行间距：固定值20磅
- 段落首行缩进2个字符
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', font_size=None, bold=False):
    """设置字体"""
    run.font.name = font_name_en
    run.font.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)


def set_paragraph_format(paragraph, line_spacing=20, first_line_indent=True):
    """设置段落格式"""
    # 固定行间距20磅
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # 首行缩进2个字符（约0.85cm，小四12pt时2个字符≈0.85cm）
    if first_line_indent:
        pf.first_line_indent = Cm(0.85)


def add_heading1(doc, text):
    """添加一级标题：三号宋体(16pt)、加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, font_size=16, bold=True)
    set_paragraph_format(p, first_line_indent=False)
    return p


def add_heading2(doc, text):
    """添加二级标题：四号宋体(14pt)、加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, font_size=14, bold=True)
    set_paragraph_format(p, first_line_indent=False)
    return p


def add_heading3(doc, text):
    """添加三级标题：小四宋体(12pt)、加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, font_size=12, bold=True)
    set_paragraph_format(p, first_line_indent=True)
    return p


def add_body_text(doc, text):
    """添加正文：小四宋体(12pt)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, font_size=12, bold=False)
    set_paragraph_format(p, first_line_indent=True)
    return p


def add_sub_item(doc, text):
    """添加子项（a) b) c)等）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, font_size=12, bold=False)
    set_paragraph_format(p, first_line_indent=True)
    return p


def parse_markdown(md_file, docx_file):
    """解析markdown并生成docx"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 逐行解析
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 跳过空行
        if not line:
            i += 1
            continue

        # 跳过分隔线
        if line == '---':
            i += 1
            continue

        # 一级标题
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_heading1(doc, text)
            i += 1
            continue

        # 二级标题
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            add_heading2(doc, text)
            i += 1
            continue

        # 三级标题
        if line.startswith('### '):
            text = line[4:].strip()
            add_heading3(doc, text)
            i += 1
            continue

        # 普通正文
        # 处理加粗文本
        if '**' in line:
            # 移除markdown加粗标记，因为标题已经处理过了
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            add_body_text(doc, clean_text)
        else:
            add_body_text(doc, line)

        i += 1

    doc.save(docx_file)
    print(f"文档已保存到: {docx_file}")


if __name__ == '__main__':
    import os
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(base_dir, '第7章_软件测试.md')
    docx_file = os.path.join(base_dir, '第7章_软件测试.docx')
    parse_markdown(md_file, docx_file)
